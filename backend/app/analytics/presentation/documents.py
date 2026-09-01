"""Render a compiled report as PDF or DOCX.

Both writers walk ``CompiledReport.blocks`` in order and handle the same closed
set of block kinds. Neither selects facts, resolves citations or decides what a
section contains — that happened once, in the compiler. What differs here is
typography and what each format can express.

Charts arrive as PNGs drawn from the same ChartSpec the Workbench displayed;
see rasterize.py for why they are drawn again rather than captured.

The PDF is the published deliverable: fixed-layout, carrying the run's figures
as rendered. The DOCX is an editable convenience copy and says so on its own
first page, because nothing can stop a reader changing a number in Word and it
should not pretend otherwise.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape

from docx import Document
from docx.document import Document as WordDocument
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    Image,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from app.analytics.presentation.compiler import EDITABLE_COPY_NOTICE
from app.analytics.presentation.document_model import CompiledMetric, CompiledReport, EvidenceEntry
from app.analytics.presentation.theme import ReportTheme

#: The appendix restates a display's own rows, which are already bounded.
MAX_APPENDIX_ROWS = 100

ChartImages = dict[str, Path]


def _cell(value: object) -> str:
    """Format one value for display. Formatting only — never a computation."""

    if isinstance(value, float):
        return f"{value:,.2f}".rstrip("0").rstrip(".")
    if isinstance(value, int) and not isinstance(value, bool):
        return f"{value:,}"
    return "" if value is None else str(value)


def _subtitle(report: CompiledReport) -> str:
    parts = [report.subtitle] if report.subtitle else []
    if report.displayed_period:
        parts.append(report.displayed_period)
    parts.append(f"Generated {report.generated_at.strftime('%d %B %Y')}")
    return "  ·  ".join(parts)


def _footer_text(report: CompiledReport) -> str:
    return f"{report.title} · Run {report.run_id}"


def _metric_note(metric: CompiledMetric) -> str:
    """Say where a figure came from, in the detail the run actually recorded."""

    parts = []
    if metric.source_query_ids:
        parts.append(", ".join(metric.source_query_ids))
    if metric.source_column:
        parts.append(f"column {metric.source_column}")
    if metric.row_selector and metric.row_selector.fields:
        parts.append(f"row {metric.row_selector.describe()}")
    return " · ".join(parts)


def _figure_caption(block: Any) -> str:
    parts = []
    if block.source_query_ids:
        parts.append(f"Based on {', '.join(block.source_query_ids)}")
    if getattr(block, "period", None):
        parts.append(block.period)
    if getattr(block, "caption", None):
        parts.append(block.caption)
    return " · ".join(parts)


def _evidence_rows(entry: EvidenceEntry) -> list[list[str]]:
    """Everything the runtime recorded about one query, as a readable grid."""

    rows = [["Item", "Detail"]]
    if entry.executed_at:
        rows.append(["Executed", entry.executed_at.strftime("%d %B %Y, %H:%M UTC")])
    if entry.period:
        rows.append(["Reporting period", entry.period])
    if entry.parameters:
        rows.append(["Parameters", ", ".join(f"{key}={value}" for key, value in entry.parameters.items())])
    if entry.tables_consulted:
        rows.append(["Tables consulted", ", ".join(entry.tables_consulted)])
    if entry.returned_columns:
        rows.append(["Returned columns", ", ".join(entry.returned_columns)])
    if entry.row_count is not None:
        rows.append(["Rows returned", f"{entry.row_count:,}{' (truncated)' if entry.truncated else ''}"])
    if entry.used_by:
        rows.append(["Used by", ", ".join(entry.used_by)])
    return rows


# ---------------------------------------------------------------------- DOCX


def write_docx(report: CompiledReport, images: ChartImages, path: Path,
               theme: ReportTheme | None = None) -> Path:
    """Write the report as an editable Word document."""

    return _DocxWriter(report, images, theme).write(path)


class _DocxWriter:
    def __init__(self, report: CompiledReport, images: ChartImages, theme: ReportTheme | None) -> None:
        self.report = report
        self.images = images
        self.theme = theme or ReportTheme()

    def write(self, path: Path) -> Path:
        document = Document()
        self._page_setup(document)
        self._running_titles(document)

        for block in self.report.blocks:
            if block.heading and block.kind not in {"cover", "page_break"}:
                document.add_heading(block.heading, level=1)
            getattr(self, f"_{block.kind}")(document, block)

        path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(path))
        return path

    # -- page furniture

    def _page_setup(self, document: WordDocument) -> None:
        normal = document.styles["Normal"]
        normal.font.name = self.theme.fonts.docx_body
        normal.font.size = Pt(self.theme.fonts.body_size)
        if self.report.orientation == "landscape":
            for section in document.sections:
                width, height = section.page_width, section.page_height
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width, section.page_height = max(width, height), min(width, height)

    def _running_titles(self, document: WordDocument) -> None:
        """A repeating header, and a footer carrying an automatic page number."""

        section = document.sections[0]
        header = section.header.paragraphs[0]
        header.text = self.report.title
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._muted(header, self.theme.fonts.caption_size)

        footer = section.footer.paragraphs[0]
        footer.text = f"{_footer_text(self.report)}  ·  Page "
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Word computes the number itself; a literal would be wrong the moment
        # the document reflows in an editor.
        _page_number_field(footer)
        self._muted(footer, self.theme.fonts.caption_size)

    def _muted(self, paragraph: Any, size: float) -> None:
        for run in paragraph.runs:
            run.font.size = Pt(size)
            run.font.color.rgb = RGBColor.from_string(self.theme.palette.muted.lstrip("#").upper())

    # -- blocks

    def _cover(self, document: WordDocument, block: Any) -> None:
        heading = document.add_paragraph()
        run = heading.add_run(block.title)
        run.bold = True
        run.font.size = Pt(self.theme.fonts.title_size)
        subtitle = document.add_paragraph()
        subtitle_run = subtitle.add_run(_subtitle(self.report))
        subtitle_run.font.size = Pt(self.theme.fonts.caption_size + 1.5)
        subtitle_run.font.color.rgb = RGBColor.from_string(self.theme.palette.muted.lstrip("#").upper())
        self._note(document, EDITABLE_COPY_NOTICE)

    def _page_break(self, document: WordDocument, block: Any) -> None:
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def _narrative(self, document: WordDocument, block: Any) -> None:
        if block.warning:
            self._warning(document, block.warning)
        if not block.lines:
            document.add_paragraph(block.empty_message)
            return
        for line in block.lines:
            if line.kind == "heading":
                document.add_heading(line.text, level=min(line.level + 1, 4))
            elif line.kind == "bullet":
                document.add_paragraph(line.text, style="List Bullet")
            elif line.kind == "number":
                document.add_paragraph(line.text, style="List Number")
            else:
                document.add_paragraph(line.text)

    def _metrics(self, document: WordDocument, block: Any) -> None:
        if not block.metrics:
            document.add_paragraph(block.empty_message)
            return
        if self.theme.metrics_style == "cards":
            self._cards(document, block.metrics)
        else:
            self._grid(document, [["Metric", "Value", "Change"]] + [
                [metric.label, metric.display_value, metric.change or ""] for metric in block.metrics
            ])
        for metric in block.metrics:
            if note := _metric_note(metric):
                self._note(document, f"{metric.label}: {note}")

    def _chart(self, document: WordDocument, block: Any) -> None:
        document.add_paragraph().add_run(f"{block.figure_label}. {block.title}").bold = True
        image = self.images.get(block.chart_id)
        if image is None:
            document.add_paragraph("This chart could not be rendered for print.")
            return
        document.add_picture(str(image), width=Inches(6.2))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._note(document, _figure_caption(block))

    def _table(self, document: WordDocument, block: Any) -> None:
        if not block.data.columns:
            document.add_paragraph(block.empty_message)
            return
        document.add_paragraph().add_run(f"{block.figure_label}. {block.title}").bold = True
        self._grid(document, [block.data.columns] + [
            [_cell(row.get(column)) for column in block.data.columns] for row in block.data.rows
        ])
        if block.data.is_truncated:
            self._note(document, f"Showing {len(block.data.rows)} of {block.data.total_row_count} rows.")
        self._note(document, _figure_caption(block))

    def _caveats(self, document: WordDocument, block: Any) -> None:
        if not block.stated:
            document.add_paragraph(block.empty_message)
        for caveat in block.stated:
            document.add_paragraph(caveat, style="List Bullet")
        for notice in block.system_notices:
            self._note(document, notice)

    def _scope(self, document: WordDocument, block: Any) -> None:
        self._grid(document, [["Item", "Detail"]] + [[row.label, row.value] for row in block.rows])

    def _evidence(self, document: WordDocument, block: Any) -> None:
        if not block.entries:
            document.add_paragraph(block.empty_message)
            return
        for entry in block.entries:
            document.add_paragraph().add_run(f"{entry.query_id} — {entry.description}").bold = True
            self._grid(document, _evidence_rows(entry))
            for label, rows in entry.displayed_rows.items():
                if not rows.columns:
                    continue
                document.add_paragraph(f"Rows shown in {label}")
                self._grid(document, [rows.columns] + [
                    [_cell(row.get(column)) for column in rows.columns]
                    for row in rows.rows[:MAX_APPENDIX_ROWS]
                ])
        if block.note:
            self._note(document, block.note)

    # -- helpers

    def _cards(self, document: WordDocument, metrics: list[CompiledMetric]) -> None:
        """The same row of cards the PDF sets, in what Word can express.

        Layout only: label, display value and change, each exactly as recorded.
        """

        table = document.add_table(rows=1, cols=len(metrics))
        table.style = "Light Grid Accent 1"
        muted = RGBColor.from_string(self.theme.palette.muted.lstrip("#").upper())
        accent = RGBColor.from_string(self.theme.palette.accent.lstrip("#").upper())
        for cell, metric in zip(table.rows[0].cells, metrics, strict=True):
            label = cell.paragraphs[0]
            label_run = label.add_run(metric.label.upper())
            label_run.font.size = Pt(self.theme.fonts.caption_size)
            label_run.font.color.rgb = muted
            value_run = cell.add_paragraph().add_run(metric.display_value)
            value_run.bold = True
            value_run.font.size = Pt(self.theme.fonts.heading_size + 3)
            value_run.font.color.rgb = accent
            if metric.change:
                change_run = cell.add_paragraph().add_run(metric.change)
                change_run.font.size = Pt(self.theme.fonts.caption_size)

    def _grid(self, document: WordDocument, rows: list[list[str]]) -> None:
        table = document.add_table(rows=1, cols=len(rows[0]))
        table.style = "Light Grid Accent 1"
        for index, heading in enumerate(rows[0]):
            table.rows[0].cells[index].paragraphs[0].add_run(heading).bold = True
        _repeat_header_row(table.rows[0])
        for row in rows[1:]:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = value

    def _warning(self, document: WordDocument, text: str) -> None:
        """Set a mismatch apart from the prose it qualifies, so it is not skimmed."""

        run = document.add_paragraph().add_run(text)
        run.bold = True
        run.font.size = Pt(self.theme.fonts.body_size)
        run.font.color.rgb = RGBColor.from_string(
            self.theme.palette.warning.lstrip("#").upper()
        )

    def _note(self, document: WordDocument, text: str) -> None:
        run = document.add_paragraph().add_run(text)
        run.font.size = Pt(self.theme.fonts.caption_size)
        run.italic = True
        run.font.color.rgb = RGBColor.from_string(self.theme.palette.muted.lstrip("#").upper())


def _page_number_field(paragraph: Any) -> None:
    """Insert a Word PAGE field, which Word evaluates when it paginates."""

    run = paragraph.add_run()
    begin, instruction, end = OxmlElement("w:fldChar"), OxmlElement("w:instrText"), OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, end):
        run._r.append(element)


def _repeat_header_row(row: Any) -> None:
    """Mark a row as a heading so Word repeats it across page breaks."""

    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


# ----------------------------------------------------------------------- PDF


def write_pdf(report: CompiledReport, images: ChartImages, path: Path,
              theme: ReportTheme | None = None) -> Path:
    """Write the report as a fixed-layout PDF."""

    return _PdfWriter(report, images, theme).write(path)


class _PdfWriter:
    def __init__(self, report: CompiledReport, images: ChartImages, theme: ReportTheme | None) -> None:
        self.report = report
        self.images = images
        self.theme = theme or ReportTheme()
        self.styles = self._styles()

    def _styles(self) -> dict[str, ParagraphStyle]:
        base = getSampleStyleSheet()["Normal"]
        fonts, palette = self.theme.fonts, self.theme.palette
        ink = colors.HexColor(palette.ink)
        return {
            "title": ParagraphStyle("title", parent=base, fontName=fonts.pdf_bold,
                                    fontSize=fonts.title_size, leading=fonts.title_size * 1.2,
                                    spaceAfter=4, textColor=ink),
            "subtitle": ParagraphStyle("subtitle", parent=base, fontName=fonts.pdf_body,
                                       fontSize=fonts.caption_size + 1, leading=12,
                                       textColor=colors.HexColor(palette.muted), spaceAfter=18),
            # keepWithNext moves a heading onto the next page rather than
            # leaving it stranded at the foot of one with its content overleaf.
            "section": ParagraphStyle("section", parent=base, fontName=fonts.pdf_bold,
                                      fontSize=fonts.heading_size, leading=fonts.heading_size * 1.3,
                                      spaceBefore=16, spaceAfter=self.theme.spacing.heading_gap,
                                      textColor=colors.HexColor(palette.accent),
                                      keepWithNext=True),
            "sub": ParagraphStyle("sub", parent=base, fontName=fonts.pdf_bold,
                                  fontSize=fonts.body_size + 0.5, leading=14, spaceBefore=9,
                                  spaceAfter=3, textColor=ink, keepWithNext=True),
            "body": ParagraphStyle("body", parent=base, fontName=fonts.pdf_body,
                                   fontSize=fonts.body_size, leading=fonts.body_size * 1.5,
                                   spaceAfter=6, textColor=ink, alignment=TA_LEFT),
            "cell": ParagraphStyle("cell", parent=base, fontName=fonts.pdf_body,
                                   fontSize=fonts.caption_size + 0.5,
                                   leading=(fonts.caption_size + 0.5) * 1.35, textColor=ink),
            "warning": ParagraphStyle("warning", parent=base, fontName=fonts.pdf_bold,
                                      fontSize=fonts.body_size, leading=fonts.body_size * 1.45,
                                      textColor=colors.HexColor(palette.warning),
                                      borderColor=colors.HexColor(palette.warning),
                                      borderWidth=0.6, borderPadding=7,
                                      spaceBefore=4, spaceAfter=12),
            "caption": ParagraphStyle("caption", parent=base, fontName=fonts.pdf_body,
                                      fontSize=fonts.caption_size, leading=10,
                                      textColor=colors.HexColor(palette.muted), spaceAfter=10),
        }

    def write(self, path: Path) -> Path:
        path.parent.mkdir(parents=True, exist_ok=True)
        story: list[Any] = []
        for block in self.report.blocks:
            flowables = getattr(self, f"_{block.kind}")(block)
            if block.heading and block.kind not in {"cover", "page_break"}:
                heading = Paragraph(_text(block.heading), self.styles["section"])
                # A chart is one unsplittable image, so keepWithNext cannot pull
                # its heading along: the pair has to be bundled, or the section
                # title is left stranded at the foot of the previous page. Every
                # other block can split, so its heading only needs to lead.
                headed = [heading, *flowables]
                flowables = [KeepTogether(headed)] if block.kind == "chart" else headed
            story.extend(flowables)

        margin = self.theme.spacing.margin
        page_size = landscape(A4) if self.report.orientation == "landscape" else A4
        SimpleDocTemplate(
            str(path), pagesize=page_size, title=self.report.title, author="AI Data Analyst",
            subject=f"Run {self.report.run_id}",
            leftMargin=margin, rightMargin=margin,
            topMargin=margin * 0.75, bottomMargin=margin * 0.75,
        ).build(story, onFirstPage=self._furniture, onLaterPages=self._furniture)
        return path

    def _furniture(self, canvas: Any, document: Any) -> None:
        """Draw the running header, footer and page number on every page."""

        canvas.saveState()
        width, height = document.pagesize
        margin = self.theme.spacing.margin
        canvas.setFont(self.theme.fonts.pdf_body, self.theme.fonts.caption_size)
        canvas.setFillColor(colors.HexColor(self.theme.palette.muted))
        canvas.setStrokeColor(colors.HexColor(self.theme.palette.rule))
        canvas.setLineWidth(0.5)
        canvas.drawRightString(width - margin, height - margin * 0.5, self.report.title)
        canvas.line(margin, height - margin * 0.6, width - margin, height - margin * 0.6)
        canvas.line(margin, margin * 0.62, width - margin, margin * 0.62)
        canvas.drawString(margin, margin * 0.4, _footer_text(self.report))
        canvas.drawRightString(width - margin, margin * 0.4, f"Page {document.page}")
        canvas.restoreState()

    # -- blocks

    def _cover(self, block: Any) -> list[Any]:
        return [
            Paragraph(_text(block.title), self.styles["title"]),
            Paragraph(_text(_subtitle(self.report)), self.styles["subtitle"]),
        ]

    def _page_break(self, block: Any) -> list[Any]:
        return [PageBreak()]

    def _narrative(self, block: Any) -> list[Any]:
        opening: list[Any] = []
        if block.warning:
            opening.append(Paragraph(_text(block.warning), self.styles["warning"]))
        if not block.lines:
            return opening + [Paragraph(_text(block.empty_message), self.styles["body"])]
        story: list[Any] = list(opening)
        pending: list[str] = []
        pending_kind = "bullet"
        for line in block.lines:
            if line.kind in {"bullet", "number"}:
                if pending and line.kind != pending_kind:
                    story.extend(self._list(pending, pending_kind))
                    pending = []
                pending_kind = line.kind
                pending.append(_text(line.text))
                continue
            story.extend(self._list(pending, pending_kind))
            pending = []
            story.append(Paragraph(_text(line.text),
                                   self.styles["sub" if line.kind == "heading" else "body"]))
        story.extend(self._list(pending, pending_kind))
        return story

    def _metrics(self, block: Any) -> list[Any]:
        if not block.metrics:
            return [Paragraph(_text(block.empty_message), self.styles["body"])]
        if self.theme.metrics_style == "cards":
            story: list[Any] = [self._cards(block.metrics), Spacer(1, self.theme.spacing.block_gap)]
        else:
            rows = [["Metric", "Value", "Change"]] + [
                [metric.label, metric.display_value, metric.change or ""] for metric in block.metrics
            ]
            story = [self._table_flowable(rows), Spacer(1, self.theme.spacing.block_gap)]
        notes = [f"{metric.label}: {note}" for metric in block.metrics if (note := _metric_note(metric))]
        if notes:
            story.append(Paragraph(_text(" · ".join(notes)), self.styles["caption"]))
        return story

    def _chart(self, block: Any) -> list[Any]:
        title = Paragraph(_text(f"{block.figure_label}. {block.title}"), self.styles["sub"])
        image = self.images.get(block.chart_id)
        if image is None:
            return [title, Paragraph("This chart could not be rendered for print.", self.styles["body"])]
        wide = self.report.orientation == "landscape"
        # Kept together so a figure caption never sits on the page after its
        # picture, and a heading never dangles at the foot of a page alone.
        return [KeepTogether([
            title,
            Image(str(image), width=(21 if wide else 16) * cm,
                  height=(8.5 if wide else 8) * cm, kind="proportional"),
            Paragraph(_text(_figure_caption(block)), self.styles["caption"]),
        ])]

    def _table(self, block: Any) -> list[Any]:
        if not block.data.columns:
            return [Paragraph(_text(block.empty_message), self.styles["body"])]
        rows = [block.data.columns] + [
            [_cell(row.get(column)) for column in block.data.columns] for row in block.data.rows
        ]
        note = _figure_caption(block)
        if block.data.is_truncated:
            note = f"Showing {len(block.data.rows)} of {block.data.total_row_count} rows. {note}"
        return [
            Paragraph(_text(f"{block.figure_label}. {block.title}"), self.styles["sub"]),
            self._table_flowable(rows),
            Paragraph(_text(note), self.styles["caption"]),
        ]

    def _caveats(self, block: Any) -> list[Any]:
        story: list[Any] = []
        if block.stated:
            story.extend(self._list([_text(caveat) for caveat in block.stated], "bullet"))
        else:
            story.append(Paragraph(_text(block.empty_message), self.styles["body"]))
        story.extend(Paragraph(_text(notice), self.styles["caption"]) for notice in block.system_notices)
        return story

    def _scope(self, block: Any) -> list[Any]:
        rows = [["Item", "Detail"]] + [[row.label, row.value] for row in block.rows]
        return [self._table_flowable(rows), Spacer(1, self.theme.spacing.block_gap)]

    def _evidence(self, block: Any) -> list[Any]:
        if not block.entries:
            return [Paragraph(_text(block.empty_message), self.styles["body"])]
        story: list[Any] = []
        for entry in block.entries:
            story.append(Paragraph(_text(f"{entry.query_id} — {entry.description}"), self.styles["sub"]))
            story.append(self._table_flowable(_evidence_rows(entry)))
            for label, rows in entry.displayed_rows.items():
                if not rows.columns:
                    continue
                story.append(Paragraph(_text(f"Rows shown in {label}"), self.styles["caption"]))
                story.append(self._table_flowable([rows.columns] + [
                    [_cell(row.get(column)) for column in rows.columns]
                    for row in rows.rows[:MAX_APPENDIX_ROWS]
                ]))
            story.append(Spacer(1, self.theme.spacing.block_gap))
        if block.note:
            story.append(Paragraph(_text(block.note), self.styles["caption"]))
        return story

    # -- helpers

    def _list(self, items: list[str], bullet: str = "bullet") -> list[Any]:
        if not items:
            return []
        # `start` numbers an ordered list; passing it to a bulleted one replaces
        # the glyph with that literal, which is how every bullet became a "1".
        options = ({"bulletType": "1", "start": 1} if bullet == "number"
                   else {"bulletType": "bullet", "bulletFontSize": self.theme.fonts.body_size,
                         "bulletOffsetY": -1})
        return [
            ListFlowable([ListItem(Paragraph(item, self.styles["body"]), leftIndent=14) for item in items],
                         leftIndent=12, **options),
            Spacer(1, 6),
        ]

    def _cards(self, metrics: list[CompiledMetric]) -> Table:
        """Set headline figures as a row of cards, the way a dashboard is read.

        Layout only: each card prints the label, the display value and the
        change exactly as the run recorded them.
        """

        palette, fonts = self.theme.palette, self.theme.fonts
        label = ParagraphStyle("card_label", fontName=fonts.pdf_body, fontSize=fonts.caption_size,
                               leading=fonts.caption_size * 1.4,
                               textColor=colors.HexColor(palette.muted))
        value = ParagraphStyle("card_value", fontName=fonts.pdf_bold, fontSize=fonts.heading_size + 5,
                               leading=(fonts.heading_size + 5) * 1.15,
                               textColor=colors.HexColor(palette.accent), spaceBefore=3)
        change = ParagraphStyle("card_change", fontName=fonts.pdf_body, fontSize=fonts.caption_size,
                                leading=fonts.caption_size * 1.4,
                                textColor=colors.HexColor(palette.ink), spaceBefore=2)
        # Each card is a stacked cell; the row of them is one table.
        cells = [[
            Paragraph(_text(metric.label.upper()), label),
            Paragraph(_text(metric.display_value), value),
            Paragraph(_text(metric.change or " "), change),
        ] for metric in metrics]
        width = landscape(A4)[0] if self.report.orientation == "landscape" else A4[0]
        column = (width - 2 * self.theme.spacing.margin) / max(len(cells), 1)
        table = Table([cells], colWidths=[column] * len(cells), hAlign="LEFT")
        table.setStyle(TableStyle([
            ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor(self.theme.palette.rule)),
            ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor(self.theme.palette.rule)),
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(self.theme.palette.table_header)),
            ("LEFTPADDING", (0, 0), (-1, -1), 10),
            ("RIGHTPADDING", (0, 0), (-1, -1), 10),
            ("TOPPADDING", (0, 0), (-1, -1), 9),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 9),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        return table

    def _table_flowable(self, rows: list[list[str]]) -> Table:
        """Build a table whose heading repeats on every page it spans."""

        palette = self.theme.palette
        table = Table(
            [[Paragraph(_text(_cell(cell)), self.styles["cell"]) for cell in row] for row in rows],
            repeatRows=1, hAlign="LEFT",
        )
        style = [
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(palette.table_header)),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]
        if self.theme.table_style == "grid":
            style.append(("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor(palette.rule)))
        else:
            style.append(("LINEBELOW", (0, 0), (-1, -1), 0.4, colors.HexColor(palette.rule)))
        table.setStyle(TableStyle(style))
        return table


def _text(value: str) -> str:
    """Render prose as characters rather than as PDF markup.

    reportlab's Paragraph parses a small HTML-like vocabulary and quietly drops
    tags it does not know, so an unescaped ``<script>`` would disappear and
    change what a sentence says. Escaping first keeps the printed wording equal
    to the compiled wording.
    """

    return escape(value)
