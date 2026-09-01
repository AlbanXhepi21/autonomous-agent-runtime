"""The Limitations section, in both formats a report is published as.

Two kinds of statement share the section and must stay distinguishable: what
the analysis said about itself, and what the runtime says about any report of
this shape. Neither is written at publishing time — the first was stored with
the run, the second is a constant.
"""

from datetime import datetime, timezone
from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfReader

from app.analytics.presentation.documents import write_docx, write_pdf
from app.analytics.presentation.compiler import CITATION_NOTICE, compile_report
from app.analytics.presentation.templates import ReportTemplateRegistry
from app.contracts.answers import AnswerSource

CAVEATS = [
    "Refund timing may differ from order timing.",
    "August 2026 is a partial month, so the total understates the period.",
]


def _source() -> AnswerSource:
    return AnswerSource(id="query_003", run_id="run-1", label="Revenue by category",
                        referenced_tables=["orders"], row_count=2)


def _report(caveats: list[str] | None = None, sources: list[AnswerSource] | None = None,
            template: str = "monthly_business_review"):
    return compile_report(
        template=ReportTemplateRegistry().get(template), run_id="run-1",
        answer="Revenue grew 18%.", charts=[],
        sources=[_source()] if sources is None else sources,
        generated_at=datetime(2026, 8, 31, tzinfo=timezone.utc), caveats=caveats,
    )


def _caveats_block(report):
    """The single compiled Limitations block, or None when the shape omits it."""

    blocks = report.blocks_of("caveats")
    return blocks[0] if blocks else None


def _pdf_text(path: Path) -> str:
    return "\n".join(page.extract_text() for page in PdfReader(str(path)).pages)


def _docx_text(path: Path) -> str:
    return "\n".join(paragraph.text for paragraph in Document(str(path)).paragraphs)


def test_a_run_that_stated_nothing_still_carries_the_system_notice() -> None:
    report = _report(caveats=None)

    block = _caveats_block(report)
    assert block is not None
    assert block.stated == []
    assert block.system_notices == [CITATION_NOTICE]


def test_the_two_kinds_of_limitation_are_held_apart() -> None:
    report = _report(caveats=CAVEATS)

    block = _caveats_block(report)
    assert block is not None
    assert block.stated == CAVEATS
    assert block.system_notices == [CITATION_NOTICE]
    # A model caveat never leaks into the system list, or the reverse.
    assert CITATION_NOTICE not in block.stated


def test_a_report_citing_nothing_omits_a_notice_about_citations() -> None:
    """There is no citation for the notice to qualify."""

    report = _report(caveats=None, sources=[])

    block = _caveats_block(report)
    assert block is not None
    assert block.system_notices == []


@pytest.mark.parametrize("writer,read", [(write_pdf, _pdf_text), (write_docx, _docx_text)])
def test_stated_limitations_are_printed_in_both_formats(writer, read, tmp_path: Path) -> None:
    path = writer(_report(caveats=CAVEATS), {}, tmp_path / "report.out")

    text = read(path).replace("\n", " ")

    for caveat in CAVEATS:
        assert caveat in text
    assert " ".join(CITATION_NOTICE.split()) in " ".join(text.split())


@pytest.mark.parametrize("writer,read", [(write_pdf, _pdf_text), (write_docx, _docx_text)])
def test_the_notice_prints_even_when_the_analysis_stated_nothing(writer, read, tmp_path: Path) -> None:
    path = writer(_report(caveats=None), {}, tmp_path / "report.out")

    text = " ".join(read(path).split())

    assert "This analysis stated no limitations." in text
    assert " ".join(CITATION_NOTICE.split()) in text


@pytest.mark.parametrize("writer,read", [(write_pdf, _pdf_text), (write_docx, _docx_text)])
def test_both_formats_print_the_same_limitations(writer, read, tmp_path: Path) -> None:
    """The two formats differ in typography, not in what they say."""

    path = writer(_report(caveats=CAVEATS), {}, tmp_path / "report.out")
    text = " ".join(read(path).split())

    assert [caveat in text for caveat in CAVEATS] == [True, True]


@pytest.mark.parametrize("writer,read", [(write_pdf, _pdf_text), (write_docx, _docx_text)])
def test_markup_in_a_caveat_is_printed_not_interpreted(writer, read, tmp_path: Path) -> None:
    """reportlab parses an HTML-like vocabulary and drops tags it does not know.

    Unescaped, ``<script>…</script>`` vanished from the page and the limitation
    read differently from what the analysis stated.
    """

    hostile = "<script>alert(1)</script> Sample of 12 orders & falling."

    path = writer(_report(caveats=[hostile]), {}, tmp_path / "report.out")
    text = " ".join(read(path).split())

    assert hostile in text


def test_every_shipped_template_can_state_limitations() -> None:
    """A published report always cites queries, so each shape must say what that proves."""

    for template in ReportTemplateRegistry().list_templates():
        kinds = [block.kind for block in template.blocks]
        assert "caveats" in kinds, f"{template.name} cannot print its limitations"


def test_limitations_appear_before_the_evidence_they_qualify(tmp_path: Path) -> None:
    path = write_docx(_report(caveats=CAVEATS), {}, tmp_path / "report.docx")

    headings = [item.text for item in Document(str(path)).paragraphs
                if item.style.name.startswith("Heading")]

    assert headings.index("Limitations") < headings.index("Evidence Appendix")
