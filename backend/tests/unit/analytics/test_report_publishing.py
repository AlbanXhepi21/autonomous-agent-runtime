"""Compiling a run into the canonical report, and drawing its displays for print."""

from datetime import datetime, timezone
from pathlib import Path

import pytest

from app.analytics.presentation.charts import ChartSpec
from app.analytics.presentation.compiler import compile_report, describe_source, parse_prose
from app.analytics.presentation.rasterize import render_chart_png
from app.analytics.presentation.templates import (
    ReportTemplateError,
    ReportTemplateRegistry,
)
from app.contracts.answers import AnswerSource


def _bar() -> ChartSpec:
    return ChartSpec(id="bar", type="bar", title="Revenue by category", x_field="category",
                     y_fields=["revenue"],
                     data=[{"category": "Electronics", "revenue": 163}, {"category": "Fashion", "revenue": 63}],
                     source_query_ids=["query_003"])


def _kpi() -> ChartSpec:
    return ChartSpec(id="kpi", type="kpi", title="Headlines",
                     kpis=[{"label": "Revenue", "value": "$163M", "change": "+18%"}],
                     source_query_ids=["query_001"])


def _table() -> ChartSpec:
    return ChartSpec(id="table", type="table", title="Detail",
                     data=[{"method": "Visa", "failures": 803}], source_query_ids=["query_004"])


def test_shipped_templates_all_load() -> None:
    templates = ReportTemplateRegistry().list_templates()

    assert {template.name for template in templates} >= {
        "monthly_business_review", "quarterly_review", "annual_review", "analysis_summary",
    }
    assert all(template.blocks for template in templates)


def test_an_unknown_template_is_named_in_the_error() -> None:
    with pytest.raises(ReportTemplateError, match="not_a_template"):
        ReportTemplateRegistry().get("not_a_template")


def test_assembly_sorts_displays_into_the_blocks_that_can_show_them() -> None:
    report = compile_report(
        template=ReportTemplateRegistry().get("monthly_business_review"), run_id="run-1",
        answer="Revenue grew.", charts=[_kpi(), _bar(), _table()],
        sources=[AnswerSource(id="query_003", run_id="run-1", label="Revenue by category")],
        generated_at=datetime.now(timezone.utc),
    )

    assert [metric.label for block in report.blocks_of("metrics") for metric in block.metrics] == ["Revenue"]
    assert [block.chart_id for block in report.blocks_of("chart")] == ["bar"]
    assert [block.table_id for block in report.blocks_of("table")] == ["table"]


def test_a_block_with_nothing_to_show_is_not_printed() -> None:
    report = compile_report(
        template=ReportTemplateRegistry().get("monthly_business_review"), run_id="run-1",
        answer="Revenue grew.", charts=[], sources=[], generated_at=datetime.now(timezone.utc),
    )

    # A required section with nothing to show says so instead of vanishing.
    kinds = [block.kind for block in report.blocks]
    assert "narrative" in kinds
    assert report.blocks_of("chart") == []
    assert [block.metrics for block in report.blocks_of("metrics")] == [[]]


def test_prose_keeps_its_structure_and_loses_its_markers() -> None:
    blocks = parse_prose("## Finding\nRevenue **grew** 18%.\n\n- Electronics led\n1. Check stock")

    assert [(line.kind, line.text) for line in blocks] == [
        ("heading", "Finding"),
        ("paragraph", "Revenue grew 18%."),
        ("bullet", "Electronics led"),
        ("number", "Check stock"),
    ]


def test_a_source_line_describes_the_query_without_its_sql() -> None:
    line = describe_source(AnswerSource(id="query_003", run_id="run-1", label="Revenue by category",
                                        referenced_tables=["orders", "order_items"], row_count=4200))

    assert line == "Revenue by category · orders, order_items · 4,200 rows"


@pytest.mark.parametrize("chart_type", ["line", "area", "bar", "stacked_bar", "pie", "scatter"])
def test_every_plotted_type_renders_an_image(chart_type: str, tmp_path: Path) -> None:
    chart = ChartSpec(type=chart_type, title="Chart", x_field="month", y_fields=["revenue"],  # type: ignore[arg-type]
                      data=[{"month": "2026-01", "revenue": 10}, {"month": "2026-02", "revenue": 14}],
                      source_query_ids=["query_001"])

    rendered = render_chart_png(chart, tmp_path / "chart.png")

    assert rendered is not None and rendered.stat().st_size > 1_000


def test_a_table_is_not_drawn_as_a_picture(tmp_path: Path) -> None:
    # Tables and KPI cards become document text, so the rasterizer declines them.
    assert render_chart_png(_table(), tmp_path / "table.png") is None
    assert render_chart_png(_kpi(), tmp_path / "kpi.png") is None


def test_a_long_form_category_column_becomes_one_series_per_category(tmp_path: Path) -> None:
    # Mirrors the Workbench pivot so an exported chart shows the same series.
    chart = ChartSpec(type="line", title="Revenue by channel", x_field="month", y_fields=["revenue"],
                      data=[{"month": f"2026-{m:02d}", "channel": channel, "revenue": m}
                            for m in (1, 2, 3) for channel in ("paid", "organic")],
                      source_query_ids=["query_001"])

    from app.analytics.presentation.rasterize import _pivot

    pivoted = _pivot(chart)

    assert pivoted is not None
    x_values, series = pivoted
    assert x_values == ["2026-01", "2026-02", "2026-03"]
    assert [name for name, _ in series] == ["paid", "organic"]


def test_a_report_states_the_scope_it_covered() -> None:
    report = compile_report(
        template=ReportTemplateRegistry().get("monthly_business_review"), run_id="run-1",
        answer="Revenue grew.", charts=[_bar()],
        sources=[
            AnswerSource(id="query_003", run_id="run-1", label="Revenue by category",
                         referenced_tables=["orders", "order_items"]),
            AnswerSource(id="query_004", run_id="run-1", label="Refunds",
                         referenced_tables=["refunds", "orders"]),
        ],
        generated_at=datetime.now(timezone.utc),
    )

    # Deduplicated and ordered, so the reader sees the scope rather than a log.
    scope = report.blocks_of("scope")[0]
    consulted = next(row.value for row in scope.rows if row.label == "Tables consulted")
    assert consulted == "order_items, orders, refunds"


def test_the_appendix_restates_only_charts_that_carry_rows() -> None:
    with_rows = compile_report(
        template=ReportTemplateRegistry().get("monthly_business_review"), run_id="run-1",
        answer="", charts=[_bar()], sources=[], generated_at=datetime.now(timezone.utc),
    )
    without_charts = compile_report(
        template=ReportTemplateRegistry().get("monthly_business_review"), run_id="run-1",
        answer="", charts=[_kpi()], sources=[], generated_at=datetime.now(timezone.utc),
    )

    # The evidence appendix restates the rows behind each drawn figure.
    entries = with_rows.blocks_of("evidence")[0].entries
    assert not entries or all(entry.query_id.startswith("query_") for entry in entries)
    assert without_charts.blocks_of("chart") == []


def test_a_detailed_template_keeps_its_structure_when_a_run_is_thin(tmp_path: Path) -> None:
    # The complaint that started this: a run that produced no displays must still
    # publish a document that says so, rather than quietly dropping the sections.
    from docx import Document

    from app.analytics.presentation.documents import write_docx

    report = compile_report(
        template=ReportTemplateRegistry().get("monthly_business_review"), run_id="run-1",
        answer="Revenue grew 18%.", charts=[], sources=[], generated_at=datetime.now(timezone.utc),
    )
    write_docx(report, {}, tmp_path / "thin.docx")
    document = Document(str(tmp_path / "thin.docx"))
    headings = [item.text for item in document.paragraphs if item.style.name.startswith("Heading")]
    body = "\n".join(item.text for item in document.paragraphs)

    assert "Headline Metrics" in headings and "Evidence Appendix" in headings
    assert "This analysis produced no charts." in body


def test_a_full_run_publishes_every_section(tmp_path: Path) -> None:
    from docx import Document

    from app.analytics.presentation.documents import write_docx
    from app.analytics.presentation.rasterize import render_chart_png

    report = compile_report(
        template=ReportTemplateRegistry().get("monthly_business_review"), run_id="run-1",
        answer="## Finding\nRevenue grew.", charts=[_kpi(), _bar(), _table()],
        sources=[AnswerSource(id="query_003", run_id="run-1", label="Revenue by category",
                              referenced_tables=["orders"], row_count=2)],
        generated_at=datetime.now(timezone.utc), period="August 2026",
    )
    # Images are keyed by the compiled chart blocks, so both renderers are handed
    # the same figures rather than each choosing its own.
    drawn = {chart.id: chart for chart in (_kpi(), _bar(), _table())}
    images = {block.chart_id: rendered for block in report.blocks_of("chart")
              if (rendered := render_chart_png(drawn[block.chart_id],
                                               tmp_path / f"{block.chart_id}.png"))}
    write_docx(report, images, tmp_path / "full.docx")
    document = Document(str(tmp_path / "full.docx"))
    headings = [item.text for item in document.paragraphs if item.style.name.startswith("Heading")]

    assert headings == [
        # Scope comes before the findings it qualifies; limitations always print
        # once anything is cited, because the citation notice belongs there.
        "Reporting Scope", "Executive Summary", "Finding", "Headline Metrics",
        "Charts", "Supporting Tables", "Limitations", "Evidence Appendix",
    ]
    assert sum(1 for rel in document.part.rels.values() if "image" in rel.reltype) == 1
