"""The canonical compiled report, and what the two renderers make of it.

The point of compiling once is that PDF and DOCX cannot disagree about a fact.
These tests assert that by reading the generated files back — never by comparing
bytes, which would only prove that two runs of the same code produce the same
output, not that the output says the right thing.
"""

from datetime import datetime, timezone
from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfReader

from app.analytics.presentation.charts import ChartSpec
from app.analytics.presentation.compiler import CITATION_NOTICE, compile_report
from app.analytics.presentation.document_model import CompiledReport
from app.analytics.presentation.documents import write_docx, write_pdf
from app.analytics.presentation.rasterize import render_chart_png
from app.analytics.presentation.templates import ReportTemplateRegistry
from app.contracts.answers import AnswerSource

TEMPLATES = ["monthly_business_review", "quarterly_review", "annual_review",
             "executive_dashboard", "analysis_summary"]

CAVEAT = "August 2026 is a partial month, so the total understates the period."


def _kpi() -> ChartSpec:
    return ChartSpec(
        id="kpi", type="kpi", title="Headlines", source_query_ids=["query_001"],
        kpis=[
            {"label": "Revenue", "value": "$163M", "change": "+18%", "raw_value": 163000000,
             "source_column": "revenue", "row_selector": {"category": "Electronics"}},
            {"label": "Orders", "value": "12,400", "raw_value": 12400, "source_column": "orders",
             "row_selector": {"category": "Electronics"}},
        ],
    )


def _bar() -> ChartSpec:
    return ChartSpec(
        id="bar", type="bar", title="Revenue by category", x_field="category",
        y_fields=["revenue"], source_query_ids=["query_003"],
        data=[{"category": "Electronics", "revenue": 163}, {"category": "Fashion", "revenue": 63}],
    )


def _table() -> ChartSpec:
    return ChartSpec(
        id="grid", type="table", title="Payment failures", source_query_ids=["query_004"],
        data=[{"method": "Visa", "failures": 803}, {"method": "Amex", "failures": 91}],
    )


def _sources() -> list[AnswerSource]:
    return [
        AnswerSource(id="query_001", run_id="run-1", label="Headline revenue and orders",
                     referenced_tables=["orders"], columns=["revenue", "orders"], row_count=1,
                     executed_at=datetime(2026, 8, 31, 9, 15, tzinfo=timezone.utc)),
        AnswerSource(id="query_003", run_id="run-1", label="Revenue by category",
                     referenced_tables=["orders", "order_items"], columns=["category", "revenue"],
                     row_count=2, executed_at=datetime(2026, 8, 31, 9, 16, tzinfo=timezone.utc)),
        AnswerSource(id="query_004", run_id="run-1", label="Payment failures by method",
                     referenced_tables=["payments"], columns=["method", "failures"], row_count=2,
                     executed_at=datetime(2026, 8, 31, 9, 17, tzinfo=timezone.utc)),
    ]


def _compile(template: str = "monthly_business_review", **overrides) -> CompiledReport:
    arguments = {
        "template": ReportTemplateRegistry().get(template),
        "run_id": "run-1",
        "answer": "## Finding\nRevenue grew 18%.\n\n- Electronics led the quarter.",
        "charts": [_kpi(), _bar(), _table()],
        "sources": _sources(),
        "generated_at": datetime(2026, 8, 31, 12, 0, tzinfo=timezone.utc),
        "period": "August 2026",
        "caveats": [CAVEAT],
    }
    return compile_report(**{**arguments, **overrides})


def _render(report: CompiledReport, directory: Path, template: str) -> tuple[Path, Path]:
    theme = ReportTemplateRegistry().get(template).theme
    drawn = {chart.id: chart for chart in (_kpi(), _bar(), _table())}
    images = {
        block.chart_id: rendered
        for block in report.blocks_of("chart")
        if (rendered := render_chart_png(drawn[block.chart_id],
                                         directory / f"{block.chart_id}.png",
                                         palette=theme.chart_palette))
    }
    return (write_pdf(report, images, directory / "report.pdf", theme),
            write_docx(report, images, directory / "report.docx", theme))


def _pdf_text(path: Path) -> str:
    return " ".join("\n".join(page.extract_text() for page in PdfReader(str(path)).pages).split())


def _docx_text(path: Path) -> str:
    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        parts.extend(cell.text for row in table.rows for cell in row.cells)
    return " ".join(" ".join(parts).split())


# ------------------------------------------------------------- canonical model


def test_the_compiled_report_records_what_produced_it() -> None:
    report = _compile()

    assert report.run_id == "run-1"
    assert report.template_id == "monthly_business_review"
    assert report.template_version
    assert report.report_id
    assert report.analysis_period == "August 2026"
    # Without a rerun the prose and the figures describe the same period.
    assert report.displayed_period == report.analysis_period
    assert report.narrative_period_status == "current"
    assert report.narrative_warning is None


def test_blocks_are_a_closed_discriminated_set() -> None:
    report = _compile()

    kinds = {block.kind for block in report.blocks}
    assert kinds <= {"cover", "scope", "narrative", "metrics", "chart",
                     "table", "caveats", "evidence", "page_break"}
    # Round-tripping proves the discriminator resolves every block back to its
    # own type rather than to a permissive base.
    restored = CompiledReport.model_validate_json(report.model_dump_json())
    assert [block.kind for block in restored.blocks] == [block.kind for block in report.blocks]


def test_the_model_refuses_a_field_it_does_not_define() -> None:
    payload = _compile().model_dump(mode="json")
    payload["blocks"][0]["injected"] = "value"

    with pytest.raises(ValueError):
        CompiledReport.model_validate(payload)


def test_compiling_the_same_run_twice_produces_the_same_document() -> None:
    """Deterministic: the report is a function of the run and the template."""

    first = _compile(report_id="fixed")
    second = _compile(report_id="fixed")

    assert first.model_dump_json() == second.model_dump_json()


# ------------------------------------------------------------------ provenance


def test_a_headline_metric_carries_the_cell_it_was_read_from() -> None:
    metrics = [metric for block in _compile().blocks_of("metrics") for metric in block.metrics]

    revenue = next(metric for metric in metrics if metric.label == "Revenue")
    assert revenue.display_value == "$163M"
    assert revenue.raw_value == 163000000
    assert revenue.source_query_ids == ["query_001"]
    assert revenue.source_column == "revenue"
    assert revenue.row_selector is not None
    assert revenue.row_selector.fields == {"category": "Electronics"}
    assert revenue.provenance_is_complete


def test_a_chart_carries_the_exact_rows_it_was_drawn_from() -> None:
    block = _compile().blocks_of("chart")[0]

    assert block.source_query_ids == ["query_003"]
    assert block.data.columns == ["category", "revenue"]
    assert block.data.rows == [
        {"category": "Electronics", "revenue": 163},
        {"category": "Fashion", "revenue": 63},
    ]
    assert block.period == "August 2026"


def test_a_table_carries_its_displayed_rows_and_scope() -> None:
    block = _compile().blocks_of("table")[0]

    assert block.source_query_ids == ["query_004"]
    assert block.data.columns == ["method", "failures"]
    assert [row["method"] for row in block.data.rows] == ["Visa", "Amex"]
    assert block.data.total_row_count == 2
    assert block.data.is_truncated is False


def test_compiling_invents_no_value_the_run_did_not_produce() -> None:
    """Every printed number must appear somewhere in the run's own displays."""

    report = _compile()
    supplied = {"$163M", "12,400", "163", "63", "803", "91"}

    for block in report.blocks_of("metrics"):
        for metric in block.metrics:
            assert metric.display_value in supplied
    for block in report.blocks_of("chart") + report.blocks_of("table"):
        for row in block.data.rows:
            for value in row.values():
                assert str(value) in supplied or isinstance(value, str)


def test_a_metric_without_recorded_provenance_stays_incomplete() -> None:
    """Absent provenance is reported as absent, never filled in."""

    bare = ChartSpec(id="kpi", type="kpi", title="Headlines", source_query_ids=["query_001"],
                     kpis=[{"label": "Revenue", "value": "$163M"}])

    metric = _compile(charts=[bare]).blocks_of("metrics")[0].metrics[0]

    assert metric.raw_value is None
    assert metric.source_column is None
    assert metric.row_selector is None
    assert metric.provenance_is_complete is False
    # The query it came from is still known, because the display cited it.
    assert metric.source_query_ids == ["query_001"]


# -------------------------------------------------------------------- evidence


def test_the_appendix_accounts_for_every_query_the_report_uses() -> None:
    entries = _compile().blocks_of("evidence")[0].entries

    assert [entry.query_id for entry in entries] == ["query_001", "query_003", "query_004"]
    by_id = {entry.query_id: entry for entry in entries}
    revenue = by_id["query_003"]
    assert revenue.description == "Revenue by category"
    assert revenue.executed_at is not None
    assert revenue.period == "August 2026"
    assert revenue.tables_consulted == ["orders", "order_items"]
    assert revenue.returned_columns == ["category", "revenue"]
    assert revenue.row_count == 2
    assert revenue.used_by == ["Figure 1"]
    assert revenue.displayed_rows["Figure 1"].rows[0] == {"category": "Electronics", "revenue": 163}


def test_the_appendix_links_a_metric_back_to_its_query() -> None:
    entries = {entry.query_id: entry for entry in _compile().blocks_of("evidence")[0].entries}

    assert entries["query_001"].used_by == ["Revenue", "Orders"]


def test_an_unknown_citation_never_reaches_the_appendix() -> None:
    """Resolution drops what it cannot account for, upstream of compiling."""

    chart = ChartSpec(id="bar", type="bar", title="Revenue", x_field="category",
                      y_fields=["revenue"], source_query_ids=["query_999"],
                      data=[{"category": "A", "revenue": 1}, {"category": "B", "revenue": 2}])

    report = _compile(charts=[chart], sources=[_sources()[1]])
    entries = report.blocks_of("evidence")[0].entries

    # query_999 was never resolved, so it has no entry to be accounted for.
    assert [entry.query_id for entry in entries] == ["query_003"]


# ------------------------------------------------------- rendering, both formats


@pytest.mark.parametrize("template", TEMPLATES)
def test_every_template_renders_in_both_formats(template: str, tmp_path: Path) -> None:
    report = _compile(template)

    pdf, docx = _render(report, tmp_path, template)

    assert pdf.stat().st_size > 5_000
    assert docx.stat().st_size > 5_000
    assert PdfReader(str(pdf)).pages
    assert Document(str(docx)).paragraphs


@pytest.mark.parametrize("template", TEMPLATES)
def test_both_formats_state_the_same_facts(template: str, tmp_path: Path) -> None:
    """The two renderers read one compiled report; they cannot disagree."""

    report = _compile(template)
    pdf, docx = _render(report, tmp_path, template)
    pdf_text, docx_text = _pdf_text(pdf), _docx_text(docx)

    for text in (pdf_text, docx_text):
        assert report.title in text
        for metric in (metric for block in report.blocks_of("metrics") for metric in block.metrics):
            assert metric.display_value in text, f"{metric.label} missing"
        for query_id in report.cited_query_ids:
            assert query_id in text, f"{query_id} not cited"
        assert CAVEAT in text
        assert " ".join(CITATION_NOTICE.split()) in text


@pytest.mark.parametrize("template", TEMPLATES)
def test_expected_section_headings_appear(template: str, tmp_path: Path) -> None:
    report = _compile(template)
    pdf, docx = _render(report, tmp_path, template)

    expected = [block.heading for block in report.blocks if block.heading]
    pdf_text = _pdf_text(pdf)
    docx_headings = [item.text for item in Document(str(docx)).paragraphs
                     if item.style.name.startswith("Heading")]

    for heading in expected:
        assert heading in pdf_text, f"{heading} missing from the PDF"
        assert heading in docx_headings, f"{heading} is not a DOCX heading"


@pytest.mark.parametrize("template", TEMPLATES)
def test_appendix_row_labels_are_printed(template: str, tmp_path: Path) -> None:
    """The exact displayed rows must be restated where a reader can check them."""

    report = _compile(template)
    pdf, docx = _render(report, tmp_path, template)

    for text in (_pdf_text(pdf), _docx_text(docx)):
        assert "Electronics" in text
        assert "Returned columns" in text
        assert "Figure 1" in text


def test_the_dashboard_is_landscape_and_the_review_is_not(tmp_path: Path) -> None:
    dashboard = _compile("executive_dashboard")
    review = _compile("monthly_business_review")

    dashboard_pdf, dashboard_docx = _render(dashboard, tmp_path / "d", "executive_dashboard")
    review_pdf, _ = _render(review, tmp_path / "r", "monthly_business_review")

    wide = PdfReader(str(dashboard_pdf)).pages[0].mediabox
    tall = PdfReader(str(review_pdf)).pages[0].mediabox
    assert wide.width > wide.height
    assert tall.height > tall.width
    section = Document(str(dashboard_docx)).sections[0]
    assert section.page_width > section.page_height


def test_the_dashboard_keeps_its_visual_summary_before_the_detail() -> None:
    kinds = [block.kind for block in _compile("executive_dashboard").blocks]

    assert kinds.index("metrics") < kinds.index("chart") < kinds.index("page_break")
    assert kinds.index("page_break") < kinds.index("evidence")


def test_the_dashboard_caps_its_charts() -> None:
    """Up to four principal charts; a run with more shows the first four."""

    charts = [
        ChartSpec(id=f"c{index}", type="bar", title=f"Chart {index}", x_field="category",
                  y_fields=["revenue"], source_query_ids=["query_003"],
                  data=[{"category": "A", "revenue": 1}, {"category": "B", "revenue": 2}])
        for index in range(7)
    ]

    report = _compile("executive_dashboard", charts=charts)

    assert len(report.blocks_of("chart")) == 4
    assert [block.figure_label for block in report.blocks_of("chart")] == [
        "Figure 1", "Figure 2", "Figure 3", "Figure 4",
    ]


def test_an_explicit_page_break_reaches_the_pdf(tmp_path: Path) -> None:
    report = _compile("monthly_business_review")

    pdf, _ = _render(report, tmp_path, "monthly_business_review")

    assert any(block.kind == "page_break" for block in report.blocks)
    assert len(PdfReader(str(pdf)).pages) > 1


def test_the_pdf_numbers_its_pages(tmp_path: Path) -> None:
    report = _compile("monthly_business_review")

    pdf, _ = _render(report, tmp_path, "monthly_business_review")
    pages = PdfReader(str(pdf)).pages

    for number, page in enumerate(pages, start=1):
        assert f"Page {number}" in " ".join(page.extract_text().split())


def test_the_docx_carries_running_titles_and_a_page_field(tmp_path: Path) -> None:
    report = _compile("monthly_business_review")

    _, docx = _render(report, tmp_path, "monthly_business_review")
    section = Document(str(docx)).sections[0]
    footer_xml = section.footer.paragraphs[0]._p.xml

    assert report.title in section.header.paragraphs[0].text
    assert report.run_id in section.footer.paragraphs[0].text
    # Word evaluates the field itself, so the number is right after an edit.
    assert "PAGE" in footer_xml


def test_the_docx_repeats_table_headings_across_pages(tmp_path: Path) -> None:
    report = _compile("monthly_business_review")

    _, docx = _render(report, tmp_path, "monthly_business_review")
    tables = Document(str(docx)).tables

    assert tables
    assert all("tblHeader" in table.rows[0]._tr.xml for table in tables)


def test_the_docx_embeds_its_charts_as_images(tmp_path: Path) -> None:
    report = _compile("monthly_business_review")

    _, docx = _render(report, tmp_path, "monthly_business_review")
    document = Document(str(docx))

    images = sum(1 for rel in document.part.rels.values() if "image" in rel.reltype)
    assert images == len(report.blocks_of("chart"))


def test_the_docx_says_it_is_an_editable_copy(tmp_path: Path) -> None:
    """It cannot guarantee its figures once a reader opens it in Word."""

    report = _compile("monthly_business_review")

    _, docx = _render(report, tmp_path, "monthly_business_review")

    assert "editing and reuse" in _docx_text(docx)
    assert "the PDF is the published report" in _docx_text(docx)


def _libreoffice() -> str | None:
    """LibreOffice, if this environment has it. Used to render a DOCX for real."""

    import shutil

    for candidate in ("soffice", "libreoffice",
                      "/Applications/LibreOffice.app/Contents/MacOS/soffice"):
        if found := shutil.which(candidate) or (candidate if Path(candidate).is_file() else None):
            return found
    return None


@pytest.mark.skipif(_libreoffice() is None, reason="LibreOffice is not installed here")
def test_the_docx_converts_to_a_readable_pdf(tmp_path: Path) -> None:
    """Word documents are laid out by Word, so the only honest check is to render one.

    python-docx writes the XML but never paginates it; a heading hierarchy or a
    repeating table header can be well-formed and still lay out badly. Where a
    converter exists, the DOCX is turned into pages and read back.
    """

    import subprocess

    report = _compile("monthly_business_review")
    _, docx = _render(report, tmp_path, "monthly_business_review")

    subprocess.run(
        [_libreoffice(), "--headless", "--convert-to", "pdf", "--outdir", str(tmp_path), str(docx)],
        check=True, capture_output=True, timeout=180,
    )
    converted = tmp_path / f"{docx.stem}.pdf"
    assert converted.is_file(), "LibreOffice produced no PDF"

    text = _pdf_text(converted)
    assert report.title in text
    for metric in (metric for block in report.blocks_of("metrics") for metric in block.metrics):
        assert metric.display_value in text
    for query_id in report.cited_query_ids:
        assert query_id in text
    assert CAVEAT in text


def test_a_template_that_declines_a_format_is_not_rendered_in_it() -> None:
    registry = ReportTemplateRegistry()

    assert registry.get("monthly_business_review").supports("pdf")
    assert not registry.get("monthly_business_review").supports("html")
